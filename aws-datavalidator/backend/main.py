"""
AR Automation — FastAPI Backend
S3 Buckets:
  Input  : s3://claude-test-tube/emails/
  Output : s3://claude-test-tube/output/
  Approved: s3://claude-test-tube/Approved/   (JSON payload)
  Reject : s3://claude-test-tube/Reject/      (XLSX)

On Approve: payload is also POSTed to CUSTOMER_API_URL downstream.
"""

import io
import json
import logging
import os
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

import boto3
import httpx
import pandas as pd
from botocore.exceptions import ClientError, NoCredentialsError
from dotenv import load_dotenv
from fastapi import FastAPI, HTTPException, Query, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

load_dotenv()

# ── Logging ──────────────────────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s  %(levelname)-8s  %(message)s",
)
log = logging.getLogger("ar-automation")


# ── CloudWatch Logs handler ───────────────────────────────────────────────────
import queue
import threading

CW_GROUP = "ar-automation-logs"
_cw_client   = None
_cw_queues   = {}          # stream_name → queue.Queue
_cw_tokens   = {}          # stream_name → sequenceToken
_cw_lock     = threading.Lock()

def _get_cw():
    global _cw_client
    if _cw_client is None:
        _cw_client = boto3.client("logs", region_name=os.getenv("AWS_DEFAULT_REGION", "ap-south-1"))
    return _cw_client

def _ensure_stream(stream: str):
    """Create log group + stream if they don't exist (idempotent)."""
    cw = _get_cw()
    try:
        cw.create_log_group(logGroupName=CW_GROUP)
    except Exception:
        pass
    try:
        cw.create_log_stream(logGroupName=CW_GROUP, logStreamName=stream)
    except Exception:
        pass

def _sender_for(stream: str, q: queue.Queue):
    """One background sender thread per log stream."""
    def _run():
        token = None
        while True:
            try:
                first = q.get(timeout=10)
                batch = [first]
                try:
                    while len(batch) < 25:
                        batch.append(q.get_nowait())
                except queue.Empty:
                    pass
                kwargs = dict(
                    logGroupName=CW_GROUP,
                    logStreamName=stream,
                    logEvents=sorted(batch, key=lambda e: e["timestamp"]),
                )
                if token:
                    kwargs["sequenceToken"] = token
                try:
                    resp  = _get_cw().put_log_events(**kwargs)
                    token = resp.get("nextSequenceToken")
                except Exception:
                    pass
            except queue.Empty:
                pass
    t = threading.Thread(target=_run, daemon=True, name=f"cw-{stream}")
    t.start()

def _cw_log(stream: str, message: str, level: str = "INFO"):
    """
    Send a log line to a specific CloudWatch stream.
    Creates the stream + sender thread on first use (lazy).
    Non-blocking — enqueues immediately.
    """
    try:
        with _cw_lock:
            if stream not in _cw_queues:
                _ensure_stream(stream)
                q = queue.Queue(maxsize=500)
                _cw_queues[stream] = q
                _sender_for(stream, q)
            q = _cw_queues[stream]
        ts  = int(__import__("time").time() * 1000)
        fmt = f"{__import__('datetime').datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')}  {level:<8}  {message}"
        q.put_nowait({"timestamp": ts, "message": fmt})
    except Exception:
        pass  # never let logging break the app

def _setup_cloudwatch_logging():
    """
    Attach a non-blocking CloudWatch handler for general/startup logs
    using the shared 'ar-automation-api' stream.
    Per-file logging uses _cw_log(stem, message) directly.
    """
    try:
        _ensure_stream("ar-automation-api")

        class CWHandler(logging.Handler):
            def emit(self, record):
                _cw_log("ar-automation-api", self.format(record), record.levelname)

        handler = CWHandler()
        handler.setFormatter(logging.Formatter("%(asctime)s  %(levelname)-8s  %(message)s"))
        logging.getLogger("ar-automation").addHandler(handler)
        log.info("CloudWatch logging enabled → %s / ar-automation-api", CW_GROUP)
    except Exception as e:
        log.warning("CloudWatch logging not available: %s", e)


_setup_cloudwatch_logging()
REGION            = os.getenv("AWS_DEFAULT_REGION", "ap-south-1")
BUCKET            = os.getenv("S3_BUCKET", "claude-test-tube")
INPUT_PREFIX      = os.getenv("S3_INPUT_PREFIX",    "emails/")
OUTPUT_PREFIX     = os.getenv("S3_OUTPUT_PREFIX",   "output/")
APPROVED_PREFIX   = os.getenv("S3_APPROVED_PREFIX", "Approved/")
REJECT_PREFIX     = os.getenv("S3_REJECT_PREFIX",   "Reject/")
DELETED_PREFIX    = os.getenv("S3_DELETED_PREFIX",  "Deleted/")
CORS_ORIGINS      = os.getenv("CORS_ORIGINS", "http://localhost:3000").split(",")

# ── Customer downstream API via AWS API Gateway ───────────────────────────────
CUSTOMER_API_URL  = os.getenv("CUSTOMER_API_URL", "")
CUSTOMER_API_KEY  = os.getenv("CUSTOMER_API_KEY", "")

# ── FastAPI ────────────────────────────────────────────────────────────────────
app = FastAPI(
    title="AR Automation API",
    description="S3-backed AR document validation workflow",
    version="2.0.0",
)
app.add_middleware(
    CORSMiddleware,
    allow_origins=CORS_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ── S3 helpers ────────────────────────────────────────────────────────────────
def s3():
    return boto3.client("s3", region_name=REGION)


def s3_put(key: str, body: bytes, content_type: str) -> None:
    try:
        s3().put_object(Bucket=BUCKET, Key=key, Body=body, ContentType=content_type)
        log.info("PUT s3://%s/%s", BUCKET, key)
    except NoCredentialsError:
        raise HTTPException(503, "AWS credentials not configured.")
    except ClientError as e:
        raise HTTPException(502, f"S3 error: {e.response['Error']['Code']}")


def s3_get(key: str) -> bytes:
    try:
        resp = s3().get_object(Bucket=BUCKET, Key=key)
        return resp["Body"].read()
    except NoCredentialsError:
        raise HTTPException(503, "AWS credentials not configured.")
    except ClientError as e:
        code = e.response["Error"]["Code"]
        if code == "NoSuchKey":
            raise HTTPException(404, f"Key not found: {key}")
        raise HTTPException(502, f"S3 error: {code}")


def s3_list(prefix: str) -> List[Dict]:
    """List objects under a prefix, return [{key, name, size, last_modified}]."""
    try:
        paginator = s3().get_paginator("list_objects_v2")
        items = []
        for page in paginator.paginate(Bucket=BUCKET, Prefix=prefix):
            for obj in page.get("Contents", []):
                key = obj["Key"]
                name = key[len(prefix):]        # strip prefix
                if not name or name.endswith("/"):  # skip "folder" keys
                    continue
                items.append({
                    "key": key,
                    "name": name,
                    "size": obj["Size"],
                    "last_modified": obj["LastModified"].isoformat(),
                    "prefix": prefix,
                })
        return items
    except NoCredentialsError:
        raise HTTPException(503, "AWS credentials not configured.")
    except ClientError as e:
        raise HTTPException(502, f"S3 error: {e.response['Error']['Code']}")


def s3_delete(key: str) -> None:
    try:
        s3().delete_object(Bucket=BUCKET, Key=key)
        log.info("DELETE s3://%s/%s", BUCKET, key)
    except (NoCredentialsError, ClientError):
        pass  # best-effort


# ── Tiny helpers ──────────────────────────────────────────────────────────────
def _str(v) -> str:
    if v is None or (isinstance(v, float) and str(v) == 'nan'):
        return ""
    return str(v).strip()

def _float(v) -> float:
    try:
        if v is None or (isinstance(v, float) and str(v) == 'nan'):
            return 0.0
        return float(v)
    except (TypeError, ValueError):
        return 0.0

def _convert_date(date_str) -> str:
    """
    Normalise any date string to YYYY-MM-DD — the format Oracle accepts for dates.
    Handles: YYYY-MM-DD, DD/MM/YYYY, DD-MM-YYYY, ISO timestamps.
    Returns empty string if input is None/empty.
    """
    if not date_str or str(date_str).strip() in ("", "None", "nan"):
        return ""
    date_str = str(date_str).strip()
    # Strip ISO timestamp suffix: "2025-10-18T00:00:00Z" → "2025-10-18"
    if "T" in date_str:
        date_str = date_str.split("T")[0]
    # Strip time after space: "2026-06-18 16:46:13" → "2026-06-18"
    if " " in date_str and len(date_str) > 10:
        date_str = date_str.split(" ")[0]
    from datetime import datetime as _dt
    for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%m/%d/%Y"):
        try:
            return _dt.strptime(date_str, fmt).strftime("%Y-%m-%d")
        except ValueError:
            continue
    log.warning("Could not parse date '%s', sending as-is", date_str)
    return date_str


def _convert_datetime(date_str) -> str:
    """
    Normalise any datetime string to YYYY-MM-DD HH:MM:SS for Oracle.
    If no time component found, returns YYYY-MM-DD only.
    """
    if not date_str or str(date_str).strip() in ("", "None", "nan"):
        return ""
    date_str = str(date_str).strip()
    from datetime import datetime as _dt
    # Try full datetime formats
    for fmt in (
        "%Y-%m-%dT%H:%M:%SZ",       # ISO with Z
        "%Y-%m-%dT%H:%M:%S",        # ISO without Z
        "%Y-%m-%dT%H:%M:%S.%f",     # ISO with microseconds
        "%Y-%m-%dT%H:%M:%S.%fZ",    # ISO with microseconds and Z
        "%Y-%m-%d %H:%M:%S",        # Standard datetime
        "%d/%m/%Y %H:%M:%S",        # DD/MM/YYYY HH:MM:SS
        "%d-%m-%Y %H:%M:%S",        # DD-MM-YYYY HH:MM:SS
    ):
        try:
            return _dt.strptime(date_str, fmt).strftime("%Y-%m-%d %H:%M:%S")
        except ValueError:
            continue
    # Fallback: try date-only formats
    return _convert_date(date_str)


# ── Customer API caller ───────────────────────────────────────────────────────
def call_customer_api(payload: Dict) -> Dict:
    """
    POST the approved payload through AWS API Gateway to the customer API.

    API Gateway setup:
      Route:       POST /transactions
      Integration: HTTP_PROXY → https://tii.in/API_AR_Appln_Test/api/transactions
      Auth:        API Key (x-api-key header) — monitored in CloudWatch
      Logging:     Full request/response in CloudWatch Logs

    The downstream X-Api-Key is forwarded via API Gateway integration header mapping.
    Never raises — failures are logged and returned as a non-blocking result.
    """
    if not CUSTOMER_API_URL:
        log.warning("CUSTOMER_API_URL not configured — skipping downstream call")
        return {"status": "skipped", "reason": "CUSTOMER_API_URL not set"}

    # Inject transaction_type into hdr
    payload_to_send = dict(payload)
    if "hdr" in payload_to_send:
        hdr = dict(payload_to_send["hdr"])
        if not hdr.get("transaction_type"):
            hdr["transaction_type"] = "INSERT"
        payload_to_send["hdr"] = hdr

    headers = {
        "Content-Type": "application/json",
        "x-api-key": CUSTOMER_API_KEY,   # API Gateway usage plan key (header)
    }

    # Also pass as query param — some API Gateway setups require it this way
    url = CUSTOMER_API_URL
    if CUSTOMER_API_KEY and "X-Api-Key" not in url:
        sep = "&" if "?" in url else "?"
        url = f"{url}{sep}X-Api-Key={CUSTOMER_API_KEY}"

    log.info("Calling customer API via API Gateway: %s", url)
    log.info("Request headers: %s", {k: (v[:6] + "***" if k == "x-api-key" else v) for k, v in headers.items()})
    log.info("Request payload (first 300 chars): %s", json.dumps(payload_to_send)[:300])

    try:
        with httpx.Client(timeout=15.0) as client:
            resp = client.post(
                url,
                content=json.dumps(payload_to_send).encode("utf-8"),
                headers=headers,
            )
        log.info(
            "API Gateway response [%s]: %s",
            resp.status_code,
            resp.text[:500],
        )
        resp.raise_for_status()
        return {
            "status":        "success",
            "http_code":     resp.status_code,
            "response_body": resp.text,
        }
    except httpx.HTTPStatusError as e:
        log.error(
            "API Gateway HTTP error [%s]: %s",
            e.response.status_code,
            e.response.text[:500],
        )
        return {
            "status":        "api_error",
            "http_code":     e.response.status_code,
            "response_body": e.response.text,
        }
    except httpx.RequestError as e:
        log.error("API Gateway network error: %s", str(e))
        return {
            "status": "network_error",
            "reason": str(e),
        }


# ── Output filename convention ─────────────────────────────────────────────────
def output_key_for(input_name: str) -> str:
    """BAJAJMOTORS.pdf  →  output/BAJAJMOTORS_extracted.xlsx"""
    stem = Path(input_name).stem
    return f"{OUTPUT_PREFIX}{stem}_extracted.xlsx"


# ── Excel builder ──────────────────────────────────────────────────────────────
def build_excel(header: Dict, transactions: List[Dict]) -> bytes:
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        # Sheet 1 — Payment Header
        hdr_df = pd.DataFrame([{
            "VENDOR_NAME":   header.get("cust_name", ""),
            "PAY_DATE":      header.get("pay_dt", ""),
            "PAY_AMOUNT":    header.get("pay_amt", ""),
            "UTR_REFERENCE": header.get("utr", ""),
            "SOURCE":        header.get("src", ""),
            "CUSTOMER_CODE": header.get("cust_code", ""),
        }])
        hdr_df.to_excel(writer, index=False, sheet_name="Header")

        # Sheet 2 — Transactions
        dtl_df = pd.DataFrame([{
            "INVOICE_NO":  t.get("doc_no", ""),
            "DATE":        t.get("doc_dt", ""),
            "GROSS":       t.get("inv_amt", 0),
            "TDS":         t.get("tds", 0),
            "DEDUCTION":   t.get("ded", 0),
            "CASH_DISC":   t.get("disc", 0),
            "NET_AMT":     t.get("net", 0),
            "STATUS":      t.get("status", "pending"),
        } for t in transactions])
        dtl_df.to_excel(writer, index=False, sheet_name="Transactions")

        # auto-fit
        for shname in writer.sheets:
            ws = writer.sheets[shname]
            for col in ws.columns:
                w = max(len(str(c.value or "")) for c in col) + 4
                ws.column_dimensions[col[0].column_letter].width = min(w, 40)

    buf.seek(0)
    return buf.read()


# ── Pydantic models ────────────────────────────────────────────────────────────
class PaymentHeader(BaseModel):
    src:              Optional[str]   = "PDF"
    utr:              Optional[str]   = ""
    cust_name:        Optional[str]   = ""
    pay_dt:           Optional[str]   = None
    pay_amt:          Optional[float] = 0.0
    cust_code:        Optional[str]   = ""
    mail_id:          Optional[str]   = ""
    mail_dt:          Optional[str]   = None
    import_ref:       Optional[str]   = ""
    cust_payment_id:  Optional[str]   = ""


class Transaction(BaseModel):
    doc_no:  Optional[str]   = ""
    doc_dt:  Optional[str]   = ""
    inv_amt: Optional[float] = 0.0
    tds:     Optional[float] = 0.0
    ded:     Optional[float] = 0.0
    disc:    Optional[float] = 0.0
    net:     Optional[float] = 0.0
    status:  Optional[str]   = "pending"   # approved | rejected | pending


class DocumentPayload(BaseModel):
    input_key:    str                     # S3 key of original file e.g. emails/BAJAJ.pdf
    header:       PaymentHeader
    transactions: List[Transaction]


class ApproveRequest(BaseModel):
    input_key:    str
    header:       PaymentHeader
    transactions: List[Transaction]


class RejectRequest(BaseModel):
    input_key:    str
    header:       PaymentHeader
    transactions: List[Transaction]


# ── In-memory session store ────────────────────────────────────────────────────
# key: input_key  →  {header: dict, transactions: [dict]}
_store: Dict[str, Dict] = {}


# ═══════════════════════════════════════════════════════════════════════════════
#  ROUTES
# ═══════════════════════════════════════════════════════════════════════════════

# ── Health ─────────────────────────────────────────────────────────────────────
@app.get("/health", tags=["Health"])
def health():
    return {"status": "ok", "bucket": BUCKET, "version": "2.0.0"}


# ── Page 1: File list from output/ ────────────────────────────────────────────
@app.get("/api/files/output", tags=["Files"])
def list_output_files():
    """
    Returns all .xlsx files present in s3://claude-test-tube/output/
    These are the extracted files shown on the dashboard (Page 1).
    """
    items = s3_list(OUTPUT_PREFIX)
    result = []
    for it in items:
        name = it["name"]                        # e.g. BAJAJMOTORS_extracted.xlsx
        # derive company name from stem
        stem = Path(name).stem                   # BAJAJMOTORS_extracted
        company = stem.replace("_extracted", "").replace("_", " ")
        # check approved / rejected status from other buckets
        result.append({
            **it,
            "company": company,
            "output_name": name,
            "status": "pending",  # will be enriched by /api/files/status
        })
    return {"files": result, "count": len(result)}


@app.get("/api/files/approved", tags=["Files"])
def list_approved_files():
    """Returns all JSON files in s3://claude-test-tube/Approved/"""
    return {"files": s3_list(APPROVED_PREFIX)}


@app.get("/api/files/rejected", tags=["Files"])
def list_rejected_files():
    """Returns all XLSX files in s3://claude-test-tube/Reject/"""
    return {"files": s3_list(REJECT_PREFIX)}


@app.get("/api/files/deleted", tags=["Files"])
def list_deleted_files():
    """Returns all files in s3://claude-test-tube/Deleted/"""
    return {"files": s3_list(DELETED_PREFIX)}


@app.post("/api/file/delete", tags=["Files"])
def soft_delete_file(key: str = Query(..., description="S3 key to move to Deleted/")):
    """
    Soft-delete: move file from output/ to Deleted/ prefix.
    Also removes from Approved/ and Reject/ if present.
    """
    name = Path(key).name
    stem = Path(key).stem.replace("_extracted", "")
    dest_key = f"{DELETED_PREFIX}{name}"

    try:
        # Copy to Deleted/
        s3().copy_object(
            Bucket=BUCKET,
            CopySource={"Bucket": BUCKET, "Key": key},
            Key=dest_key,
        )
        # Delete from output/
        s3_delete(key)
        # Also clean up Approved/Reject copies
        s3_delete(f"{APPROVED_PREFIX}{stem}.json")
        s3_delete(f"{REJECT_PREFIX}{stem}_rejected.xlsx")
        # Remove from in-memory cache
        _store.pop(key, None)
        log.info("Soft-deleted %s → %s", key, dest_key)
        _cw_log(stem, f"DELETED | moved to {dest_key}")
        return {"status": "success", "deleted_key": dest_key}
    except NoCredentialsError:
        raise HTTPException(503, "AWS credentials not configured.")
    except ClientError as e:
        raise HTTPException(502, f"S3 error: {e.response['Error']['Code']}")


@app.get("/api/files/dashboard", tags=["Files"])
def dashboard_summary():
    """
    Single call for the Dashboard page.
    Returns counts + file lists for output, approved, rejected.
    """
    try:
        output   = s3_list(OUTPUT_PREFIX)
        approved = s3_list(APPROVED_PREFIX)
        rejected = s3_list(REJECT_PREFIX)
        deleted  = s3_list(DELETED_PREFIX)
    except HTTPException as exc:
        # demo mode — return empty lists
        if exc.status_code == 503:
            return {
                "output":   [], "approved": [], "rejected": [], "deleted": [],
                "counts":   {"output": 0, "approved": 0, "rejected": 0, "deleted": 0},
                "demo_mode": True,
            }
        raise

    # Enrich output list with approval/rejection status.
    # Build lookup: clean_stem → {status, last_modified}
    # If a file somehow exists in both folders, the most-recently modified one wins.
    status_map: Dict[str, Dict] = {}

    for f in approved:
        stem = Path(f["name"]).stem.lower()              # "mahavir"
        lm   = f.get("last_modified", "")
        existing = status_map.get(stem)
        if not existing or lm > existing["last_modified"]:
            status_map[stem] = {"status": "approved", "last_modified": lm}

    for f in rejected:
        stem = Path(f["name"]).stem.lower().replace("_rejected", "")  # "bajajmotors"
        lm   = f.get("last_modified", "")
        existing = status_map.get(stem)
        if not existing or lm > existing["last_modified"]:
            status_map[stem] = {"status": "rejected", "last_modified": lm}

    enriched = []
    # Build vendor name lookup from approved JSON files (they contain cust_name)
    vendor_names: Dict[str, str] = {}
    for f in approved:
        stem = Path(f["name"]).stem.lower()
        try:
            raw = s3_get(f["key"])
            data = json.loads(raw)
            name = data.get("hdr", {}).get("cust_name", "")
            if name:
                vendor_names[stem] = name
        except Exception:
            pass

    for f in output:
        raw_stem   = Path(f["name"]).stem               # MAHAVIR_extracted
        clean_stem = raw_stem.replace("_extracted", "") # MAHAVIR
        company    = clean_stem.replace("_", " ")
        source_type = "PDF"  # default source type
        mail_id    = ""      # default mail_id
        # Priority 1: read from Excel directly (no cache)
        cached = None
        if cached and cached.get("header", {}).get("cust_name"):
            company = cached["header"]["cust_name"]
        if cached and cached.get("header", {}).get("src"):
            source_type = cached["header"]["src"]
        if cached and cached.get("header", {}).get("mail_id"):
            mail_id = cached["header"]["mail_id"]
        # Priority 2: approved JSON (has vendor name from previous approval)
        elif clean_stem.lower() in vendor_names:
            company = vendor_names[clean_stem.lower()]
        # Priority 3: read vendor name from Excel (quick — first row only)
        else:
            try:
                raw = s3_get(f["key"])
                xl = pd.ExcelFile(io.BytesIO(raw))
                sheets = xl.sheet_names
                if "Remittance" in sheets:
                    rem = xl.parse("Remittance", nrows=1)
                    rem.columns = [str(c).strip() for c in rem.columns]
                    if not rem.empty and "CUSTOMER_NAME" in rem.columns:
                        name = _str(rem.iloc[0].get("CUSTOMER_NAME"))
                        if name:
                            company = name
                    if not rem.empty and "SOURCE_TYPE" in rem.columns:
                        src_val = _str(rem.iloc[0].get("SOURCE_TYPE"))
                        if src_val:
                            source_type = src_val
                    if not rem.empty and "MAIL_ID" in rem.columns:
                        mid = _str(rem.iloc[0].get("MAIL_ID"))
                        if mid:
                            mail_id = mid
                elif "Header" in sheets:
                    hdr = xl.parse("Header", nrows=1)
                    if not hdr.empty and "VENDOR_NAME" in hdr.columns:
                        name = _str(hdr.iloc[0].get("VENDOR_NAME"))
                        if name:
                            company = name
                    if not hdr.empty and "SOURCE" in hdr.columns:
                        src_val = _str(hdr.iloc[0].get("SOURCE"))
                        if src_val:
                            source_type = src_val
                    if not hdr.empty and "MAIL_ID" in hdr.columns:
                        mid = _str(hdr.iloc[0].get("MAIL_ID"))
                        if mid:
                            mail_id = mid
                else:
                    df = xl.parse(sheets[0], nrows=1)
                    df.columns = [str(c).strip() for c in df.columns]
                    for col_name in ["CUSTOMER_NAME", "VENDOR_NAME"]:
                        if col_name in df.columns and not df.empty:
                            name = _str(df.iloc[0].get(col_name))
                            if name:
                                company = name
                                break
                    for col_name in ["SOURCE_TYPE", "SOURCE", "src"]:
                        if col_name in df.columns and not df.empty:
                            src_val = _str(df.iloc[0].get(col_name))
                            if src_val:
                                source_type = src_val
                                break
                    for col_name in ["MAIL_ID", "mail_id"]:
                        if col_name in df.columns and not df.empty:
                            mid = _str(df.iloc[0].get(col_name))
                            if mid:
                                mail_id = mid
                                break
            except Exception:
                pass
        entry      = status_map.get(clean_stem.lower())
        file_status = entry["status"] if entry else "pending"
        enriched.append({**f, "company": company, "status": file_status, "source_type": source_type, "mail_id": mail_id})

    return {
        "output":   enriched,
        "approved": approved,
        "rejected": rejected,
        "deleted":  deleted,
        "counts": {
            "output":   len(output),
            "approved": len(approved),
            "rejected": len(rejected),
            "deleted":  len(deleted),
            "pending":  sum(1 for f in enriched if f["status"] == "pending"),
        },
    }


# ── Cache bust ────────────────────────────────────────────────────────────────
@app.delete("/api/file/cache", tags=["Validation"])
def clear_cache(key: str = Query(None)):
    """Clear in-memory parse cache. Pass ?key= for single file, or clear all."""
    if key:
        _store.pop(key, None)
        return {"cleared": key}
    _store.clear()
    return {"cleared": "all"}


# ── Page 2: Load a specific output file for validation ───────────────────────
@app.get("/api/file/load", tags=["Validation"])
def load_output_file(key: str = Query(..., description="S3 key of output XLSX")):
    """
    Read an _extracted.xlsx from output/ and return header + transactions
    so the frontend can render the split-pane validation view.
    """
    stem = Path(key).stem.replace("_extracted", "")

    # Check in-memory store first — DISABLED (always read fresh from S3)
    # if key in _store:
    #     _cw_log(stem, f"LOAD (cached) {key}")
    #     return {"input_key": key, **_store[key]}

    _cw_log(stem, f"LOAD started → reading from S3: {key}")

    try:
        raw = s3_get(key)
    except HTTPException as exc:
        if exc.status_code in (503, 404):
            _cw_log(stem, f"LOAD demo mode — key not found: {key}", "WARNING")
            return _demo_payload(key)
        raise

    try:
        xl = pd.ExcelFile(io.BytesIO(raw))
        sheets = xl.sheet_names
        log.info("Sheets in %s: %s", key, sheets)
        _cw_log(stem, f"LOAD sheets detected: {sheets}")

        header = {}
        transactions = []

        # ── Format A: our own output (Header + Transactions sheets) ───────────
        if "Header" in sheets and "Transactions" in sheets:
            hdr_df = xl.parse("Header")
            dtl_df = xl.parse("Transactions")
            if not hdr_df.empty:
                r = hdr_df.iloc[0]
                header = {
                    "cust_name":   _str(r.get("VENDOR_NAME")),
                    "pay_dt":      _str(r.get("PAY_DATE")),
                    "pay_amt":     _float(r.get("PAY_AMOUNT")),
                    "utr":         _str(r.get("UTR_REFERENCE")),
                    "src":         _str(r.get("SOURCE", "PDF")),
                    "cust_code":   _str(r.get("CUSTOMER_CODE")),
                    "mail_id":     _str(r.get("MAIL_ID", "")),
                    "mail_dt":     _str(r.get("MAIL_RECEIVED_DATE", "")) or None,
                    "import_ref":  _str(r.get("IMPORT_REFERENCE", "")),
                    "cust_payment_id": _str(r.get("CUST_PAYMENT_ID", "")),
                }
            for _, r in dtl_df.iterrows():
                transactions.append({
                    "doc_no":  _str(r.get("INVOICE_NO")),
                    "doc_dt":  _str(r.get("DATE")),
                    "inv_amt": _float(r.get("GROSS")),
                    "tds":     _float(r.get("TDS")),
                    "ded":     _float(r.get("DEDUCTION")),
                    "disc":    _float(r.get("CASH_DISC")),
                    "net":     _float(r.get("NET_AMT")),
                    "status":  _str(r.get("STATUS", "pending")),
                })

        # ── Format B: pipeline output (Remittance + Raw_JSON sheets) ──────────
        elif "Remittance" in sheets:
            rem = xl.parse("Remittance")
            rem.columns = [str(c).strip() for c in rem.columns]
            log.info("Remittance columns: %s", list(rem.columns))

            # ALWAYS read transactions from the Remittance sheet (it has the
            # final processed/normalized data — correct DOCUMENT_NUMBER, etc.)
            if not rem.empty:
                r0 = rem.iloc[0]
                utr_val = _str(r0.get("UTR_REFERENCE_NUMBER")) or _str(r0.get("UTR_REFERENCE_NUA")) or _str(r0.get("UTR_REFERENCE_NUM"))
                header = {
                    "cust_name":   _str(r0.get("CUSTOMER_NAME")),
                    "pay_dt":      _str(r0.get("PAYMENT_DATE")),
                    "pay_amt":     _float(r0.get("PAYMENT_AMOUNT")),
                    "utr":         utr_val,
                    "src":         _str(r0.get("SOURCE_TYPE", "PDF")),
                    "cust_code":   _str(r0.get("KOD_CUST_CODE")),
                    "mail_id":     _str(r0.get("MAIL_ID")),
                    "mail_dt":     _str(r0.get("MAIL_RECEIVED_DATE")) or None,
                    "import_ref":  _str(r0.get("IMPORT_REFERENCE", "")),
                    "cust_payment_id": _str(r0.get("CUST_PAYMENT_ID", "")),
                }
                for _, r in rem.iterrows():
                    transactions.append({
                        "doc_no":  _str(r.get("DOCUMENT_NUMBER")),
                        "doc_dt":  _str(r.get("DOCUMENT_DATE")),
                        "inv_amt": _float(r.get("INVOICE_AMOUNT")),
                        "tds":     _float(r.get("TDS_AMOUNT")),
                        "ded":     _float(r.get("DEDUCTION_AMOUNT")),
                        "disc":    _float(r.get("CASH_DISCOUNT")),
                        "net":     _float(r.get("AMOUNT")),
                        "status":  "pending",
                    })

            # Enrich header from Raw_JSON only for fields missing in Remittance
            if "Raw_JSON" in sheets:
                try:
                    rj = xl.parse("Raw_JSON")
                    if not rj.empty and "full_json" in rj.columns:
                        raw_json = json.loads(rj["full_json"].iloc[0])
                        vouchers = raw_json.get("vouchers", [])
                        if vouchers:
                            h = vouchers[0].get("header", {})
                            # Only fill header fields that are still empty
                            if not header.get("cust_name"):
                                header["cust_name"] = _str(h.get("CUSTOMER_NAME"))
                            if not header.get("pay_amt"):
                                header["pay_amt"] = _float(h.get("AMOUNT") or h.get("_computed_payment_amount") or h.get("PAYMENT_AMOUNT"))
                            if not header.get("utr"):
                                header["utr"] = _str(h.get("UTR_REFERENCE_NUMBER") or h.get("UTR_REFERENCE_NUA") or h.get("UTR_REFERENCE_NUM"))
                            if not header.get("pay_dt"):
                                header["pay_dt"] = _str(h.get("PAYMENT_DATE"))
                            if not header.get("cust_code"):
                                header["cust_code"] = _str(h.get("KOD_CUST_CODE"))
                            if not header.get("mail_id"):
                                header["mail_id"] = _str(h.get("MAIL_ID"))
                            if not header.get("mail_dt"):
                                header["mail_dt"] = _str(h.get("MAIL_RECEIVED_DATE")) or None
                except Exception as je:
                    log.warning("Raw_JSON parse failed: %s", je)

        # ── Format C: unknown — read first sheet generically ─────────────────
        else:
            df = xl.parse(sheets[0])
            # Strip whitespace from all column names
            df.columns = [str(c).strip() for c in df.columns]
            log.info("Format C columns in %s: %s", key, list(df.columns))
            col = lambda *names: next((n for n in names if n in df.columns), None)
            utr_col  = col("UTR_REFERENCE_NUMBER", "UTR_REFERENCE_NUM", "UTR_REFERENCE_NUA", "UTR", "utr")
            name_col = col("CUSTOMER_NAME", "VENDOR_NAME", "cust_name")
            dt_col   = col("PAYMENT_DATE", "PAY_DATE", "pay_dt")
            amt_col  = col("PAYMENT_AMOUNT", "PAY_AMOUNT", "pay_amt", "AMOUNT")
            doc_col  = col("DOCUMENT_NUMBER", "INVOICE_NO", "doc_no")
            ddt_col  = col("DOCUMENT_DATE", "DATE", "doc_dt")
            inv_col  = col("INVOICE_AMOUNT", "GROSS", "inv_amt")
            tds_col  = col("TDS_AMOUNT", "TDS", "tds")
            ded_col  = col("DEDUCTION_AMOUNT", "DEDUCTION", "ded")
            dsc_col  = col("CASH_DISCOUNT", "CASH_DISC", "disc")
            net_col  = col("AMOUNT", "NET_AMT", "net")
            src_col  = col("SOURCE_TYPE", "src")
            cc_col   = col("KOD_CUST_CODE", "cust_code")
            mid_col  = col("MAIL_ID", "mail_id")
            mdt_col  = col("MAIL_RECEIVED_DATE", "mail_dt")
            ref_col  = col("IMPORT_REFERENCE", "import_ref")
            cpid_col = col("CUST_PAYMENT_ID", "cust_payment_id")
            log.info("Mapped columns — utr:%s name:%s mail_id:%s mail_dt:%s",
                     utr_col, name_col, mid_col, mdt_col)

            if not df.empty:
                r0 = df.iloc[0]
                header = {
                    "cust_name":   _str(r0.get(name_col)) if name_col else "",
                    "pay_dt":      _str(r0.get(dt_col))   if dt_col   else "",
                    "pay_amt":     _float(r0.get(amt_col)) if amt_col else 0.0,
                    "utr":         _str(r0.get(utr_col))  if utr_col  else "",
                    "src":         _str(r0.get(src_col))  if src_col  else "PDF",
                    "cust_code":   _str(r0.get(cc_col))   if cc_col   else "",
                    "mail_id":     _str(r0.get(mid_col))  if mid_col  else "",
                    "mail_dt":     _str(r0.get(mdt_col))  if mdt_col  else None,
                    "import_ref":  _str(r0.get(ref_col))  if ref_col  else "",
                    "cust_payment_id": _str(r0.get(cpid_col)) if cpid_col else "",
                }
                for _, r in df.iterrows():
                    transactions.append({
                        "doc_no":  _str(r.get(doc_col))  if doc_col  else "",
                        "doc_dt":  _str(r.get(ddt_col))  if ddt_col  else "",
                        "inv_amt": _float(r.get(inv_col)) if inv_col else 0.0,
                        "tds":     _float(r.get(tds_col)) if tds_col else 0.0,
                        "ded":     _float(r.get(ded_col)) if ded_col else 0.0,
                        "disc":    _float(r.get(dsc_col)) if dsc_col else 0.0,
                        "net":     _float(r.get(net_col)) if net_col else 0.0,
                        "status":  "pending",
                    })

        payload = {"header": header, "transactions": transactions}
        # _store[key] = payload  # DISABLED — no caching, always read fresh from S3
        _cw_log(stem, f"LOAD complete — {len(transactions)} transactions | "
                      f"vendor={header.get('cust_name')} utr={header.get('utr')} "
                      f"pay_amt={header.get('pay_amt')} mail_id={header.get('mail_id')}")
        return {"input_key": key, **payload}

    except Exception as e:
        log.warning("Could not parse Excel %s: %s — returning demo", key, e)
        return _demo_payload(key)


# ── Page 2: Save in-progress state ────────────────────────────────────────────
@app.post("/api/file/save", tags=["Validation"])
def save_state(payload: DocumentPayload):
    """Auto-save current validation state to memory."""
    _store[payload.input_key] = {
        "header":       payload.header.model_dump(),
        "transactions": [t.model_dump() for t in payload.transactions],
    }
    return {"status": "saved"}


# ── Page 2: Approve → push JSON to Approved/ ──────────────────────────────────
@app.post("/api/file/approve", tags=["Validation"])
def approve_file(req: ApproveRequest):
    """
    Build the final JSON payload and push to s3://claude-test-tube/Approved/
    Also removes any existing entry from Reject/ so status is unambiguous.
    """
    stem = Path(req.input_key).stem.replace("_extracted", "")
    json_key  = f"{APPROVED_PREFIX}{stem}.json"
    old_key   = f"{REJECT_PREFIX}{stem}_rejected.xlsx"   # clean up opposite folder

    payload = {
        "hdr": {
            "src":             req.header.src or "PDF",
            "utr":             req.header.utr or "",
            "cust_name":       req.header.cust_name or "",
            "pay_dt":          _convert_date(req.header.pay_dt),
            "pay_amt":         req.header.pay_amt or 0.0,
            "cust_code":       req.header.cust_code or "",
            "mail_id":         req.header.mail_id or "",
            "mail_dt":         _convert_date(req.header.mail_dt) or None,
            "import_ref":      req.header.import_ref or "",
            "cust_payment_id": req.header.cust_payment_id or "",
        },
        "dtl": [
            {
                "doc_no":  t.doc_no,
                "doc_dt":  _convert_date(t.doc_dt),
                "inv_amt": t.inv_amt,
                "tds":     t.tds,
                "ded":     t.ded,
                "disc":    t.disc,
                "net":     t.net,
            }
            for t in req.transactions
            if t.status != "rejected"
        ],
    }

    json_bytes = json.dumps(payload, indent=2, default=str).encode("utf-8")

    demo_mode = False
    try:
        # Remove from Reject/ first (best-effort)
        s3_delete(old_key)
        # Write approved JSON to S3
        s3_put(json_key, json_bytes, "application/json")
    except HTTPException as exc:
        if exc.status_code == 503:
            log.warning("Demo mode: approve simulated for %s", req.input_key)
            demo_mode = True
        else:
            raise

    _store.pop(req.input_key, None)
    log.info("Approved %s -> %s", stem, json_key)
    _cw_log(stem, f"APPROVED | vendor={req.header.cust_name} | utr={req.header.utr} "
                  f"| pay_amt={req.header.pay_amt} | pay_dt={req.header.pay_dt} "
                  f"| transactions={len(req.transactions)} | s3={json_key}")

    # ── Call customer downstream API ──────────────────────────────────────────
    _cw_log(stem, f"API PAYLOAD | {json.dumps(payload['hdr'])}")
    api_result = call_customer_api(payload)
    log.info("Customer API result for %s: %s", stem, api_result)
    _cw_log(stem, f"CUSTOMER API result | status={api_result.get('status')} "
                  f"http_code={api_result.get('http_code')} "
                  f"response={str(api_result.get('response_body',''))[:200]}")

    return {
        "status":     "success",
        "demo_mode":  demo_mode,
        "s3_key":     f"s3://{BUCKET}/{json_key}",
        "api_result": api_result,
    }


# ── Page 2: Reject → push XLSX to Reject/ ────────────────────────────────────
@app.post("/api/file/reject", tags=["Validation"])
def reject_file(req: RejectRequest):
    """
    Build rejection XLSX and push to s3://claude-test-tube/Reject/
    Also removes any existing entry from Approved/ so status is unambiguous.
    """
    stem     = Path(req.input_key).stem.replace("_extracted", "")
    xlsx_key = f"{REJECT_PREFIX}{stem}_rejected.xlsx"
    old_key  = f"{APPROVED_PREFIX}{stem}.json"   # clean up opposite folder

    excel_bytes = build_excel(req.header.model_dump(), [t.model_dump() for t in req.transactions])

    try:
        # Remove from Approved/ first (best-effort)
        s3_delete(old_key)
        # Write to Reject/
        s3_put(xlsx_key, excel_bytes,
               "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    except HTTPException as exc:
        if exc.status_code == 503:
            log.warning("Demo mode: reject simulated for %s", req.input_key)
            return {
                "status": "success",
                "demo_mode": True,
                "s3_key": f"s3://{BUCKET}/{xlsx_key}",
            }
        raise

    _store.pop(req.input_key, None)
    log.info("Rejected %s → %s (removed approve copy if any)", stem, xlsx_key)
    _cw_log(stem, f"REJECTED | vendor={req.header.cust_name} | utr={req.header.utr} "
                  f"| pay_amt={req.header.pay_amt} | transactions={len(req.transactions)} "
                  f"| s3={xlsx_key}")

    return {
        "status": "success",
        "s3_key": f"s3://{BUCKET}/{xlsx_key}",
    }


# ── Find matching input file in emails/ ──────────────────────────────────────
@app.get("/api/file/find-input", tags=["Files"])
def find_input_file(output_key: str = Query(...)):
    """
    Given an output key like output/NISSANMOTORINDIAPVTLTD_extracted.xlsx,
    find the matching source file in emails/ by stem matching.
    Returns presigned URL if found.
    """
    stem = Path(output_key).stem.replace("_extracted", "").lower()

    try:
        items = s3_list(INPUT_PREFIX)
    except HTTPException:
        return {"found": False}

    # Try exact stem match first, then prefix match
    best = None
    for item in items:
        item_stem = Path(item["name"]).stem.lower()
        if item_stem == stem:
            best = item
            break
        if item_stem.startswith(stem) or stem.startswith(item_stem):
            best = item  # keep looking for exact

    if not best:
        return {"found": False, "stem_searched": stem}

    try:
        url = s3().generate_presigned_url(
            "get_object",
            Params={
                "Bucket": BUCKET,
                "Key": best["key"],
                "ResponseContentDisposition": "inline",
            },
            ExpiresIn=900,
        )
        return {
            "found":    True,
            "key":      best["key"],
            "name":     best["name"],
            "ext":      Path(best["name"]).suffix.lower().lstrip("."),
            "url":      url,
        }
    except (NoCredentialsError, ClientError):
        return {"found": False}


@app.get("/api/file/presign", tags=["Files"])
def presign_url(key: str = Query(...), expires: int = 900):
    """Generate a presigned GET URL for in-browser preview of any S3 object."""
    try:
        url = s3().generate_presigned_url(
            "get_object",
            Params={
                "Bucket": BUCKET,
                "Key": key,
                "ResponseContentDisposition": "inline",
            },
            ExpiresIn=expires,
        )
        return {"url": url, "expires_in": expires}
    except NoCredentialsError:
        raise HTTPException(503, "AWS credentials not configured.")
    except ClientError as e:
        raise HTTPException(502, str(e))


@app.get("/api/file/view", tags=["Files"])
def view_file(key: str = Query(..., description="S3 key to stream inline")):
    """
    Proxy-stream an S3 object through the backend so the browser renders it
    inline (same-origin). Avoids cross-origin iframe download issues with PDFs.
    """
    try:
        resp = s3().get_object(Bucket=BUCKET, Key=key)
        body = resp["Body"].read()
        # Determine content type from S3 metadata or file extension
        ct = resp.get("ContentType", "")
        if not ct or ct == "application/octet-stream":
            ext = Path(key).suffix.lower()
            ct_map = {
                ".pdf":  "application/pdf",
                ".png":  "image/png",
                ".jpg":  "image/jpeg",
                ".jpeg": "image/jpeg",
                ".gif":  "image/gif",
                ".tiff": "image/tiff",
                ".bmp":  "image/bmp",
                ".webp": "image/webp",
                ".doc":  "application/msword",
                ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ".html": "text/html",
                ".htm":  "text/html",
                ".txt":  "text/plain",
                ".xls":  "application/vnd.ms-excel",
                ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                ".csv":  "text/csv",
            }
            ct = ct_map.get(ext, "application/octet-stream")
        return StreamingResponse(
            io.BytesIO(body),
            media_type=ct,
            headers={"Content-Disposition": "inline"},
        )
    except NoCredentialsError:
        raise HTTPException(503, "AWS credentials not configured.")
    except ClientError as e:
        code = e.response["Error"]["Code"]
        if code == "NoSuchKey":
            raise HTTPException(404, f"Key not found: {key}")
        raise HTTPException(502, f"S3 error: {code}")


# ── DOC/DOCX preview → convert to HTML ───────────────────────────────────────
@app.get("/api/file/doc-preview", tags=["Files"])
def doc_preview(key: str = Query(..., description="S3 key of DOC/DOCX file")):
    """
    Convert a DOC/DOCX file from S3 to HTML for browser preview.
    Uses python-docx for .docx files.
    """
    try:
        raw = s3_get(key)
    except HTTPException:
        raise HTTPException(404, f"File not found: {key}")

    ext = Path(key).suffix.lower()

    if ext == ".docx":
        try:
            from docx import Document
            doc = Document(io.BytesIO(raw))
            html_parts = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    html_parts.append("<br/>")
                    continue
                style = para.style.name.lower() if para.style else ""
                if "heading 1" in style:
                    html_parts.append(f"<h1>{text}</h1>")
                elif "heading 2" in style:
                    html_parts.append(f"<h2>{text}</h2>")
                elif "heading 3" in style:
                    html_parts.append(f"<h3>{text}</h3>")
                else:
                    html_parts.append(f"<p>{text}</p>")

            # Also extract tables
            for table in doc.tables:
                html_parts.append("<table border='1' cellpadding='4' cellspacing='0' style='border-collapse:collapse;width:100%;margin:12px 0;'>")
                for i, row in enumerate(table.rows):
                    html_parts.append("<tr>")
                    for cell in row.cells:
                        tag = "th" if i == 0 else "td"
                        html_parts.append(f"<{tag}>{cell.text.strip()}</{tag}>")
                    html_parts.append("</tr>")
                html_parts.append("</table>")

            html = "\n".join(html_parts)
            return {"html": html, "format": "docx"}
        except Exception as e:
            log.warning("DOCX parse failed for %s: %s", key, e)
            raise HTTPException(422, f"Could not parse DOCX: {str(e)}")

    elif ext == ".doc":
        # For old .doc binary format, extract text using Word binary structure
        try:
            import re
            # Method: decode as latin-1 (preserves all bytes), then extract
            # readable text runs (sequences of printable ASCII characters)
            text = raw.decode("latin-1", errors="ignore")
            # Extract runs of readable text (at least 20 chars to avoid binary noise)
            # Word .doc files store text in specific segments
            # Look for the main text body — it's typically between specific markers
            # Alternative: extract all readable sentences
            readable_parts = re.findall(r'[\x20-\x7E]{20,}', text)

            # Filter out lines that look like binary/metadata
            clean_lines = []
            for part in readable_parts:
                # Skip lines that are mostly special chars or look like internal metadata
                alpha_ratio = sum(1 for c in part if c.isalpha() or c.isspace()) / len(part) if part else 0
                if alpha_ratio > 0.5 and not part.startswith('PK') and 'theme/theme' not in part and 'Content_Types' not in part and 'xmln' not in part:
                    clean_lines.append(part.strip())

            if clean_lines:
                # Join and format as HTML paragraphs
                # Try to find the main document text (usually the longest coherent block)
                full_text = " ".join(clean_lines)
                # Split into sentences/paragraphs at logical breaks
                paragraphs = re.split(r'(?<=[.!?])\s+(?=[A-Z])', full_text)
                html = "\n".join(f"<p>{p.strip()}</p>" for p in paragraphs if p.strip())
            else:
                html = "<p>Unable to extract text from this .doc file.</p>"

            return {"html": html, "format": "doc"}
        except Exception as e:
            log.warning("DOC parse failed for %s: %s", key, e)
            raise HTTPException(422, f"Could not parse DOC: {str(e)}")

    else:
        raise HTTPException(400, f"Unsupported file type: {ext}")


# ── Upload any file → place in output/ as _extracted.xlsx ─────────────────────
@app.post("/api/file/upload-output", tags=["Files"])
async def upload_output_file(file_key: str, content_b64: str):
    """
    Accept base64-encoded Excel content, store in output/ with _extracted suffix.
    (For programmatic uploads from ETL pipelines.)
    """
    import base64
    raw = base64.b64decode(content_b64)
    out_key = output_key_for(Path(file_key).name)
    try:
        s3_put(out_key, raw,
               "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    except HTTPException as exc:
        if exc.status_code == 503:
            return {"status": "success", "demo_mode": True, "s3_key": out_key}
        raise
    return {"status": "success", "s3_key": f"s3://{BUCKET}/{out_key}"}


# ── Demo / fallback data ───────────────────────────────────────────────────────
def _demo_payload(key: str) -> Dict[str, Any]:
    stem = Path(key).stem.replace("_extracted", "")
    return {
        "input_key": key,
        "demo_mode": True,
        "header": {
            "src": "PDF",
            "utr": "SCBLR12025101800803615",
            "cust_name": stem.upper(),
            "pay_dt": "2025-10-18",
            "pay_amt": 375469.06,
            "cust_code": "",
            "mail_id": "",
            "mail_dt": None,
        },
        "transactions": [
            {"doc_no": "4001010284952", "doc_dt": "2025-08-20", "inv_amt": 58598.40,  "tds": 46, "ded": 0, "disc": 0, "net": 58552.40,  "status": "pending"},
            {"doc_no": "4001010284953", "doc_dt": "2025-08-20", "inv_amt": 41018.88,  "tds": 33, "ded": 0, "disc": 0, "net": 40985.88,  "status": "pending"},
            {"doc_no": "4001010285360", "doc_dt": "2025-08-22", "inv_amt": 76177.92,  "tds": 60, "ded": 0, "disc": 0, "net": 76117.92,  "status": "pending"},
            {"doc_no": "4001010285359", "doc_dt": "2025-08-22", "inv_amt": 41018.88,  "tds": 33, "ded": 0, "disc": 0, "net": 40985.88,  "status": "pending"},
            {"doc_no": "4001010284490", "doc_dt": "2025-08-18", "inv_amt": 52738.56,  "tds": 42, "ded": 0, "disc": 0, "net": 52696.56,  "status": "pending"},
            {"doc_no": "4001010284489", "doc_dt": "2025-08-18", "inv_amt": 52738.56,  "tds": 42, "ded": 0, "disc": 0, "net": 52696.56,  "status": "pending"},
        ],
    }


# ── Multi-Customer: list files from multi-output/ ─────────────────────────────
MULTI_OUTPUT_PREFIX = os.getenv("S3_MULTI_OUTPUT_PREFIX", "multi-output/")
MULTI_INPUT_PREFIX  = os.getenv("S3_MULTI_INPUT_PREFIX",  "multi-input/")

@app.get("/api/files/multi-output", tags=["Multi"])
def list_multi_output_files():
    """
    Returns multi-output files split by customer.
    One file with 4 customers → 4 entries in the list, each with its own customer name.
    """
    items = s3_list(MULTI_OUTPUT_PREFIX)
    result = []
    for it in items:
        name = it["name"]
        # Try to read customer names and mail_id from the Excel
        try:
            raw = s3_get(it["key"])
            df = pd.read_excel(io.BytesIO(raw), nrows=200)
            df.columns = [str(c).strip() for c in df.columns]
            cust_col = next((c for c in df.columns if c.upper() in ("CUSTOMER_NAME", "CUST_NAME", "VENDOR_NAME")), None)
            mail_col = next((c for c in df.columns if c.upper() in ("MAIL_ID", "EMAIL", "MAIL")), None)
            mail_id  = ""
            if mail_col and not df.empty:
                mail_val = df[mail_col].dropna().unique()
                mail_id  = str(mail_val[0]).strip() if len(mail_val) > 0 else ""
            if cust_col and not df.empty:
                names = df[cust_col].dropna().unique()
                names = [str(n).strip() for n in names if str(n).strip()]
                if names:
                    for cust_name in names:
                        result.append({
                            **it,
                            "company": cust_name,
                            "status": "pending",
                            "source_type": "MULTI",
                            "mail_id": mail_id,
                        })
                    continue
        except Exception:
            pass
        # Fallback — show file stem as company
        result.append({**it, "company": Path(name).stem.replace("_", " "), "status": "pending", "source_type": "MULTI", "mail_id": ""})
    return {"files": result, "count": len(result)}


@app.get("/api/file/multi-find-input", tags=["Multi"])
def multi_find_input(output_key: str = Query(...)):
    """
    Given a multi-output key, find the matching source file in multi-input/
    by stem matching. Returns presigned URL for inline preview.
    """
    stem = Path(output_key).stem.lower()

    try:
        items = s3_list(MULTI_INPUT_PREFIX)
    except HTTPException:
        return {"found": False}

    best = None
    for item in items:
        item_stem = Path(item["name"]).stem.lower()
        if item_stem == stem:
            best = item
            break
        if item_stem.startswith(stem) or stem.startswith(item_stem):
            best = item

    if not best:
        return {"found": False, "stem_searched": stem}

    try:
        url = s3().generate_presigned_url(
            "get_object",
            Params={
                "Bucket": BUCKET,
                "Key": best["key"],
                "ResponseContentDisposition": "inline",
            },
            ExpiresIn=900,
        )
        return {
            "found": True,
            "key":   best["key"],
            "name":  best["name"],
            "ext":   Path(best["name"]).suffix.lower().lstrip("."),
            "url":   url,
        }
    except (NoCredentialsError, ClientError):
        return {"found": False}


@app.get("/api/file/multi-load", tags=["Multi"])
def load_multi_file(key: str = Query(..., description="S3 key of multi-output XLSX")):
    """
    Read a multi-customer XLSX from multi-output/ and return grouped customers.
    Each customer group has its own rows showing all columns as-is from the Excel.
    """
    stem = Path(key).stem
    try:
        raw = s3_get(key)
    except HTTPException as exc:
        if exc.status_code in (503, 404):
            return {"input_key": key, "demo_mode": True, "customers": []}
        raise

    try:
        xl   = pd.ExcelFile(io.BytesIO(raw))
        df   = xl.parse(xl.sheet_names[0])
        df.columns = [str(c).strip() for c in df.columns]

        # Normalize column names — find CUST_NO and CUSTOMER_NAME
        cust_no_col   = next((c for c in df.columns if c.upper() in ("CUST_NO", "CUST_NUMBER", "CUSTOMER_NO", "CUSTOMER_NUMBER")), None)
        cust_name_col = next((c for c in df.columns if c.upper() in ("CUSTOMER_NAME", "CUST_NAME", "VENDOR_NAME")), None)

        if not cust_no_col and not cust_name_col:
            return {"input_key": key, "demo_mode": False, "customers": [], "error": "No customer identifier column found"}

        group_col = cust_no_col or cust_name_col

        # Convert all columns to string-safe values for JSON — handle NaN safely
        import math
        def safe_str(v):
            if v is None:
                return None
            if isinstance(v, float):
                if math.isnan(v) or math.isinf(v):
                    return None
                if v == int(v):
                    return str(int(v))
                return str(v)
            return str(v)

        df = df.where(pd.notnull(df), None)
        for col in df.columns:
            df[col] = df[col].apply(safe_str)

        # Group rows by customer
        customers = []
        stem = Path(key).stem
        # Check which customers are already approved/rejected
        approved_keys = set()
        rejected_keys = set()
        try:
            for item in s3_list(APPROVED_PREFIX):
                approved_keys.add(item["name"].replace(".json", "").lower())
        except Exception:
            pass
        try:
            for item in s3_list(REJECT_PREFIX):
                rejected_keys.add(item["name"].replace("_rejected.xlsx", "").lower())
        except Exception:
            pass

        for group_key, group_df in df.groupby(group_col, sort=False):
            first = group_df.iloc[0]
            cust_name = _str(first.get(cust_name_col)) if cust_name_col else str(group_key)
            cust_no   = _str(first.get(cust_no_col))   if cust_no_col   else ""
            rows = group_df.to_dict(orient="records")

            # Check if this customer was already approved/rejected
            cust_key = f"{stem}_{cust_no}".replace("/", "_").replace(" ", "_").lower()
            if cust_key in approved_keys:
                cust_status = "approved"
                row_status  = "approved"
            elif cust_key in rejected_keys:
                cust_status = "rejected"
                row_status  = "rejected"
            else:
                cust_status = "pending"
                row_status  = "pending"

            for r in rows:
                r["_status"] = row_status
            customers.append({
                "cust_no":   cust_no,
                "cust_name": cust_name,
                "columns":   list(df.columns),
                "rows":      rows,
                "status":    cust_status,
            })

        return {"input_key": key, "demo_mode": False, "customers": customers}

    except Exception as e:
        log.warning("Multi-load failed for %s: %s", key, e)
        raise HTTPException(422, f"Could not parse multi file: {str(e)}")


class MultiApproveRequest(BaseModel):
    input_key: str
    customers: List[Dict]   # each: {cust_no, cust_name, status, rows:[...]}


class MultiCustomerApproveRequest(BaseModel):
    input_key:  str
    customer:   Dict   # single customer: {cust_no, cust_name, status, rows:[...]}


@app.post("/api/file/multi-customer-approve", tags=["Multi"])
def multi_customer_approve(req: MultiCustomerApproveRequest):
    """
    Approve a single customer from a multi-customer file.
    Builds multi-customer payload (payment vs invoice split) and sends to downstream API.
    Saves to Approved/{input_stem}_{cust_no}.json
    """
    stem     = Path(req.input_key).stem
    cust     = req.customer
    cust_no  = str(cust.get("cust_no", "")).replace("/", "_").replace(" ", "_") or "unknown"
    cust_name = cust.get("cust_name", "")
    json_key = f"{APPROVED_PREFIX}{stem}_{cust_no}.json"

    rows = cust.get("rows", [])

    # Determine source type from file extension
    input_ext = Path(req.input_key).suffix.lower()
    if input_ext in (".xlsx", ".xls"):
        src = "MULTI EXCEL"
    elif input_ext in (".txt", ".csv"):
        src = "MULTI TXT"
    else:
        src = "MULTI"

    # Get mail_id and mail_received_date from rows (same for all)
    first = rows[0] if rows else {}
    mail_id = _str(first.get("MAIL_ID"))
    mail_dt = _convert_date(first.get("MAIL_RECEIVED_DATE"))

    # Build payload: hdr = common info, dtl = ALL rows (PMT and INV)
    # PMT rows → only utr, pay_dt, pay_amt filled; doc fields null
    # INV rows → only doc_no, doc_dt, inv_amt filled; payment fields null
    # No payment info in hdr (otherwise it replicates to all rows in Oracle)

    payload = {
        "hdr": {
            "src":              src,
            "cust_name":        cust_name,
            "cust_code":        cust_no,
            "mail_id":          mail_id,
            "mail_dt":          mail_dt,
            "import_ref":       "",
            "cust_payment_id":  "",
            "transaction_type": "INSERT",
        },
        "dtl": [],
    }

    # PMT rows first, then INV rows
    for row in rows:
        if row.get("_status") == "rejected":
            continue
        class_val = str(row.get("CLASS") or "").strip().upper()
        if class_val in ("PMT", "PAYMENT"):
            payload["dtl"].append({
                "utr":     _str(row.get("TRX_NUMBER")),
                "pay_dt":  _convert_date(row.get("TXN_DATE")),
                "pay_amt": _float(row.get("OUTSTANDING_AMT") or 0),
                "doc_no":  None,
                "doc_dt":  None,
                "inv_amt": None,
                "tds":     None,
                "ded":     None,
                "disc":    None,
                "net":     _float(row.get("APPLIED_AMT") or row.get("OUTSTANDING_AMT") or 0),
            })

    for row in rows:
        if row.get("_status") == "rejected":
            continue
        class_val = str(row.get("CLASS") or "").strip().upper()
        if class_val not in ("PMT", "PAYMENT"):
            payload["dtl"].append({
                "doc_no":  _str(row.get("TRX_NUMBER")),
                "doc_dt":  _convert_date(row.get("TXN_DATE")),
                "inv_amt": _float(row.get("OUTSTANDING_AMT") or 0),
                "tds":     _float(row.get("TDS") or 0),
                "ded":     _float(row.get("DEDUCTION") or row.get("REJECTION_SHORT") or 0),
                "disc":    _float(row.get("DISCOUNT") or 0),
                "net":     _float(row.get("APPLIED_AMT") or row.get("OUTSTANDING_AMT") or 0),
                "utr":     None,
                "pay_dt":  None,
                "pay_amt": None,
            })

    try:
        s3_put(json_key, json.dumps(payload, indent=2, default=str).encode(),
               "application/json")
    except HTTPException as exc:
        if exc.status_code == 503:
            return {"status": "success", "demo_mode": True, "s3_key": json_key}
        raise

    log.info("Multi-customer approved: %s → %s", cust_name, json_key)
    api_result = call_customer_api(payload)
    log.info("Customer API result for %s: %s", cust_name, api_result)

    return {
        "status":     "success",
        "s3_key":     f"s3://{BUCKET}/{json_key}",
        "api_result": api_result,
    }


@app.post("/api/file/multi-approve", tags=["Multi"])
def multi_approve(req: MultiApproveRequest):
    """
    Approve all customers in a multi-customer file.
    Sends SEPARATE payload per customer to downstream API.
    """
    stem    = Path(req.input_key).stem
    results = []

    input_ext = Path(req.input_key).suffix.lower()
    if input_ext in (".xlsx", ".xls"):
        src = "MULTI EXCEL"
    elif input_ext in (".txt", ".csv"):
        src = "MULTI TXT"
    else:
        src = "MULTI"

    for cust in req.customers:
        if cust.get("status") == "rejected":
            results.append({"cust_no": cust.get("cust_no"), "cust_name": cust.get("cust_name"), "status": "skipped_rejected"})
            continue

        cust_no   = str(cust.get("cust_no", "")).replace("/", "_").replace(" ", "_") or "unknown"
        cust_name = cust.get("cust_name", "")
        json_key  = f"{APPROVED_PREFIX}{stem}_{cust_no}.json"
        rows      = cust.get("rows", [])
        first     = rows[0] if rows else {}
        mail_id   = _str(first.get("MAIL_ID"))
        mail_dt   = _convert_date(first.get("MAIL_RECEIVED_DATE"))

        payload = {
            "hdr": {
                "src":              src,
                "cust_name":        cust_name,
                "cust_code":        cust_no,
                "mail_id":          mail_id,
                "mail_dt":          mail_dt,
                "import_ref":       "",
                "cust_payment_id":  "",
                "transaction_type": "INSERT",
            },
            "dtl": [],
        }

        # PMT rows first, then INV rows
        for row in rows:
            if row.get("_status") == "rejected":
                continue
            class_val = str(row.get("CLASS") or "").strip().upper()
            if class_val in ("PMT", "PAYMENT"):
                payload["dtl"].append({
                    "utr":     _str(row.get("TRX_NUMBER")),
                    "pay_dt":  _convert_date(row.get("TXN_DATE")),
                    "pay_amt": _float(row.get("OUTSTANDING_AMT") or 0),
                    "doc_no":  None,
                    "doc_dt":  None,
                    "inv_amt": None,
                    "tds":     None,
                    "ded":     None,
                    "disc":    None,
                    "net":     _float(row.get("APPLIED_AMT") or row.get("OUTSTANDING_AMT") or 0),
                })

        for row in rows:
            if row.get("_status") == "rejected":
                continue
            class_val = str(row.get("CLASS") or "").strip().upper()
            if class_val not in ("PMT", "PAYMENT"):
                payload["dtl"].append({
                    "doc_no":  _str(row.get("TRX_NUMBER")),
                    "doc_dt":  _convert_date(row.get("TXN_DATE")),
                    "inv_amt": _float(row.get("OUTSTANDING_AMT") or 0),
                    "tds":     _float(row.get("TDS") or 0),
                    "ded":     _float(row.get("DEDUCTION") or row.get("REJECTION_SHORT") or 0),
                    "disc":    _float(row.get("DISCOUNT") or 0),
                    "net":     _float(row.get("APPLIED_AMT") or row.get("OUTSTANDING_AMT") or 0),
                    "utr":     None,
                    "pay_dt":  None,
                    "pay_amt": None,
                })

        try:
            s3_put(json_key, json.dumps(payload, indent=2, default=str).encode(), "application/json")
            api_result = call_customer_api(payload)
            results.append({"cust_no": cust_no, "cust_name": cust_name, "s3_key": f"s3://{BUCKET}/{json_key}", "api_result": api_result})
            log.info("Multi-customer approved: %s → %s | API: %s", cust_name, json_key, api_result.get("status"))
        except HTTPException as exc:
            if exc.status_code == 503:
                results.append({"cust_no": cust_no, "cust_name": cust_name, "s3_key": json_key, "demo_mode": True})
            else:
                raise

    _store.pop(req.input_key, None)
    log.info("Multi-approve complete: %d customers", len(results))
    return {"status": "success", "customers": results}


@app.post("/api/file/multi-reject", tags=["Multi"])
def multi_reject(req: MultiApproveRequest):
    """Save rejected multi-customer data to Reject/ as XLSX."""
    stem     = Path(req.input_key).stem
    xlsx_key = f"{REJECT_PREFIX}{stem}_rejected.xlsx"

    # Build flat Excel from all customer rows
    all_rows = []
    for cust in req.customers:
        for row in cust.get("rows", []):
            clean = {k: v for k, v in row.items() if not k.startswith("_")}
            all_rows.append(clean)

    buf = io.BytesIO()
    if all_rows:
        pd.DataFrame(all_rows).to_excel(buf, index=False)
    else:
        pd.DataFrame().to_excel(buf, index=False)
    buf.seek(0)

    try:
        s3_put(xlsx_key, buf.read(),
               "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    except HTTPException as exc:
        if exc.status_code == 503:
            return {"status": "success", "demo_mode": True, "s3_key": xlsx_key}
        raise
    _store.pop(req.input_key, None)
    log.info("Multi-rejected %s → %s", stem, xlsx_key)
    return {"status": "success", "s3_key": f"s3://{BUCKET}/{xlsx_key}"}


# ── Rejected Emails (from DynamoDB rejected_files table) ──────────────────────
REJECTED_FILES_TABLE = os.getenv("REJECTED_FILES_TABLE", "rejected_files")

@app.get("/api/rejected-emails", tags=["Rejected"])
def list_rejected_emails():
    """
    Scan the rejected_files DynamoDB table and return all entries.
    These are files rejected at Lambda level (e.g., duplicate ETag).
    """
    try:
        dynamodb_resource = boto3.resource("dynamodb", region_name=REGION)
        table = dynamodb_resource.Table(REJECTED_FILES_TABLE)
        response = table.scan()
        items = response.get("Items", [])
        # Handle pagination
        while "LastEvaluatedKey" in response:
            response = table.scan(ExclusiveStartKey=response["LastEvaluatedKey"])
            items.extend(response.get("Items", []))
        # Convert Decimal to float for JSON serialization
        from decimal import Decimal
        def convert(obj):
            if isinstance(obj, Decimal):
                return float(obj)
            return obj
        clean_items = []
        for item in items:
            clean_items.append({k: convert(v) for k, v in item.items()})
        # Sort by rejected_at descending
        clean_items.sort(key=lambda x: x.get("rejected_at", ""), reverse=True)
        return {"items": clean_items, "count": len(clean_items)}
    except Exception as e:
        log.warning("Failed to scan rejected_files table: %s", e)
        return {"items": [], "count": 0, "error": str(e)}


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
